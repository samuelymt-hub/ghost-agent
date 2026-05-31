"""Loader（文档加载）单元测试 (Req 5.1, 5.2, 5.3, 5.4, 5.5, 5.6)。

覆盖：
* 结构化标题层级与段落归属保留（md / html / 首标题前正文）。
* 纯文本与元数据归一化。
* 三类失败路径（解析失败/超时/内容为空）均终止处理并将 IngestTask 标记 FAILED。
* pdf/docx 缺少可选依赖时优雅失败。
"""

from __future__ import annotations

import time

import pytest

from ghost_agent.core import FileMeta, Loader, ParseResult, Section
from ghost_agent.models import (
    EmptyContentError,
    IngestTask,
    IngestTaskStatus,
    ParseFailedError,
    ParseTimeoutError,
)


def _task() -> IngestTask:
    return IngestTask(file_name="manual.md", file_format="md")


# --------------------------------------------------------------------------- #
# 成功路径：结构化标题层级与段落归属 (Req 5.1, 5.2, 5.4)                          #
# --------------------------------------------------------------------------- #
def test_markdown_preserves_heading_levels_order_and_paragraph_ownership() -> None:
    content = (
        "# H1 Title\n\n"
        "para under h1\n\n"
        "## H2 Title\n\n"
        "para under h2 line1\npara under h2 line2\n\n"
        "### H3 Title\n\n"
        "para under h3\n"
    )
    result = Loader().parse(content=content, file_name="a.md", file_format="md")

    assert isinstance(result, ParseResult)
    assert [(s.title, s.level) for s in result.sections] == [
        ("H1 Title", 1),
        ("H2 Title", 2),
        ("H3 Title", 3),
    ]
    # 段落归属正确
    assert result.sections[0].paragraphs == ["para under h1"]
    assert result.sections[1].paragraphs == ["para under h2 line1 para under h2 line2"]
    assert result.sections[2].paragraphs == ["para under h3"]
    # 统一纯文本包含全部标题与段落
    assert "H1 Title" in result.text
    assert "para under h3" in result.text


def test_markdown_content_before_first_heading_is_level_zero_section() -> None:
    content = "intro paragraph\n\n# First Heading\n\nbody\n"
    result = Loader().parse(content=content, file_name="a.md", file_format="md")

    assert result.sections[0].level == 0
    assert result.sections[0].title == ""
    assert result.sections[0].paragraphs == ["intro paragraph"]
    assert result.sections[1].title == "First Heading"
    assert result.sections[1].level == 1


def test_plain_text_single_level_zero_section_split_on_blank_lines() -> None:
    content = "first paragraph line\n\nsecond paragraph\n\nthird"
    result = Loader().parse(content=content, file_name="notes.txt", file_format="txt")

    assert len(result.sections) == 1
    assert result.sections[0].level == 0
    assert result.sections[0].paragraphs == [
        "first paragraph line",
        "second paragraph",
        "third",
    ]
    assert result.text.strip() != ""


def test_html_extracts_headings_and_paragraphs() -> None:
    content = (
        "<html><body>"
        "<h1>Main</h1>"
        "<p>first para</p>"
        "<h2>Sub</h2>"
        "<p>second para</p>"
        "</body></html>"
    )
    result = Loader().parse(content=content, file_name="doc.html", file_format="html")

    assert [(s.title, s.level) for s in result.sections] == [
        ("Main", 1),
        ("Sub", 2),
    ]
    assert result.sections[0].paragraphs == ["first para"]
    assert result.sections[1].paragraphs == ["second para"]


def test_bytes_content_decoded_as_utf8() -> None:
    content = "# 标题\n\n中文段落".encode("utf-8")
    result = Loader().parse(content=content, file_name="cn.md", file_format="md")
    assert result.sections[0].title == "标题"
    assert result.sections[0].paragraphs == ["中文段落"]


# --------------------------------------------------------------------------- #
# 元数据 (Req 5.4)                                                              #
# --------------------------------------------------------------------------- #
def test_meta_generates_source_file_id_when_absent() -> None:
    result = Loader().parse(content="# t\n\nbody", file_name="a.md", file_format="md")
    assert isinstance(result.meta, FileMeta)
    assert result.meta.source_file_id  # 自动生成非空
    assert result.meta.file_name == "a.md"
    assert result.meta.file_format == "md"


def test_meta_uses_provided_source_file_id_and_normalizes_format() -> None:
    result = Loader().parse(
        content="# t\n\nbody",
        file_name="a.MD",
        file_format=".MD",
        source_file_id="src-123",
    )
    assert result.meta.source_file_id == "src-123"
    assert result.meta.file_format == "md"  # 小写、去前导点


# --------------------------------------------------------------------------- #
# 失败路径：解析失败 (Req 5.3)                                                   #
# --------------------------------------------------------------------------- #
def test_unsupported_format_raises_and_marks_task_failed() -> None:
    task = _task()
    with pytest.raises(ParseFailedError):
        Loader().parse(
            content=b"data",
            file_name="x.bin",
            file_format="bin",
            ingest_task=task,
        )
    assert task.status is IngestTaskStatus.FAILED
    assert task.failure_reason and task.failure_reason.strip()


def test_non_utf8_bytes_raises_parse_failed_and_marks_task_failed() -> None:
    task = _task()
    with pytest.raises(ParseFailedError):
        Loader().parse(
            content=b"\xff\xfe\x00bad",
            file_name="bad.txt",
            file_format="txt",
            ingest_task=task,
        )
    assert task.status is IngestTaskStatus.FAILED
    assert task.failure_reason and task.failure_reason.strip()


def test_parser_raising_unexpected_error_is_classified_as_parse_failed() -> None:
    loader = Loader()

    def _boom(_content: object) -> list[Section]:
        raise RuntimeError("内部解析炸了")

    loader.register_parser("boom", _boom)
    task = IngestTask(file_name="b.boom", file_format="boom")
    with pytest.raises(ParseFailedError):
        loader.parse(
            content="x", file_name="b.boom", file_format="boom", ingest_task=task
        )
    assert task.status is IngestTaskStatus.FAILED


# --------------------------------------------------------------------------- #
# 失败路径：内容为空 (Req 5.6)                                                   #
# --------------------------------------------------------------------------- #
@pytest.mark.parametrize("content", ["", "   \n\t  \n"])
def test_empty_or_whitespace_content_raises_empty_content_and_marks_failed(
    content: str,
) -> None:
    task = IngestTask(file_name="empty.txt", file_format="txt")
    with pytest.raises(EmptyContentError):
        Loader().parse(
            content=content,
            file_name="empty.txt",
            file_format="txt",
            ingest_task=task,
        )
    assert task.status is IngestTaskStatus.FAILED
    assert task.failure_reason and task.failure_reason.strip()


# --------------------------------------------------------------------------- #
# 失败路径：解析超时 (Req 5.5)                                                   #
# --------------------------------------------------------------------------- #
def test_parse_timeout_raises_and_marks_task_failed() -> None:
    loader = Loader(parse_timeout_seconds=0.05)

    def _slow(_content: object) -> list[Section]:
        time.sleep(1.0)
        return [Section(title="", level=0, paragraphs=["late"])]

    loader.register_parser("slow", _slow)
    task = IngestTask(file_name="s.slow", file_format="slow")
    with pytest.raises(ParseTimeoutError):
        loader.parse(
            content="x", file_name="s.slow", file_format="slow", ingest_task=task
        )
    assert task.status is IngestTaskStatus.FAILED
    assert task.failure_reason and task.failure_reason.strip()


# --------------------------------------------------------------------------- #
# pdf/docx 缺少可选依赖时优雅失败                                                #
# --------------------------------------------------------------------------- #
def test_pdf_without_optional_lib_raises_parse_failed() -> None:
    pytest.importorskip  # noqa: B018 - 占位，保持可读性
    try:
        import pypdf  # type: ignore[import-not-found]  # noqa: F401

        pytest.skip("pypdf 已安装，跳过缺依赖路径校验")
    except ImportError:
        pass
    task = IngestTask(file_name="d.pdf", file_format="pdf")
    with pytest.raises(ParseFailedError):
        Loader().parse(
            content=b"%PDF-1.4 fake",
            file_name="d.pdf",
            file_format="pdf",
            ingest_task=task,
        )
    assert task.status is IngestTaskStatus.FAILED


def test_docx_without_optional_lib_raises_parse_failed() -> None:
    try:
        import docx  # type: ignore[import-not-found]  # noqa: F401

        pytest.skip("python-docx 已安装，跳过缺依赖路径校验")
    except ImportError:
        pass
    task = IngestTask(file_name="d.docx", file_format="docx")
    with pytest.raises(ParseFailedError):
        Loader().parse(
            content=b"PK fake docx",
            file_name="d.docx",
            file_format="docx",
            ingest_task=task,
        )
    assert task.status is IngestTaskStatus.FAILED


# --------------------------------------------------------------------------- #
# 不传 ingest_task 时也应正常抛错（不应因 None 而崩溃）                           #
# --------------------------------------------------------------------------- #
def test_failure_without_ingest_task_still_raises() -> None:
    with pytest.raises(EmptyContentError):
        Loader().parse(content="   ", file_name="e.txt", file_format="txt")
