"""Loader（加载器）：将受支持类型的原始文档解析为统一纯文本表示 (Req 5)。

本模块负责知识库入库管线的第一环：把上传/同步的文件解析为
:class:`ParseResult`（统一纯文本 + 结构化标题层级 + 文件元数据），并在解析失
败、解析超时或内容为空时终止处理、不移交 Transformer，同时把关联的
:class:`~ghost_agent.models.ingest_task.IngestTask` 标记为 ``FAILED``。

设计要点：

* **零额外第三方依赖**：txt / md / html 使用纯标准库解析（``re``、
  ``html.parser``）。pdf / docx 通过惰性可选导入实现，缺少依赖时优雅地抛出
  :class:`ParseFailedError`，因此模块在任何环境下都能干净导入。
* **可插拔解析器注册表**：内置解析器按文件格式注册，调用方可通过
  :meth:`Loader.register_parser` 注入自定义解析器（也便于测试注入慢解析器以
  触发超时路径）。
* **跨平台解析超时**：基于 ``concurrent.futures.ThreadPoolExecutor`` 的线程级
  超时守卫，避免 ``signal`` 在 Windows 上的不可靠行为 (Req 5.5)。

对应需求：5.1（解析为统一纯文本）、5.2（保留标题层级与段落归属）、5.3（解析
失败终止并标记 FAILED）、5.4（成功时附带来源文件标识/文件名/文件格式元数据）、
5.5（解析超时）、5.6（内容为空）。
"""

from __future__ import annotations

import re
import uuid
from concurrent.futures import ThreadPoolExecutor
from concurrent.futures import TimeoutError as FuturesTimeoutError
from html.parser import HTMLParser
from io import BytesIO
from typing import Any, Callable

from pydantic import BaseModel, ConfigDict, Field

from ghost_agent.config import get_settings
from ghost_agent.models.errors import (
    EmptyContentError,
    ParseFailedError,
    ParseTimeoutError,
)
from ghost_agent.models.ingest_task import IngestTask, IngestTaskStatus

__all__ = [
    "Section",
    "FileMeta",
    "ParseResult",
    "Loader",
    "DEFAULT_PARSE_TIMEOUT_SECONDS",
]

#: 默认解析超时（秒）。可在构造 :class:`Loader` 时通过
#: ``parse_timeout_seconds`` 覆盖。刻意不写入全局 ``Settings``，保持本组件自洽。
DEFAULT_PARSE_TIMEOUT_SECONDS: float = 30.0

#: 解析器签名：接收原始内容（bytes 或 str），返回结构化 Section 列表。
ParserFn = Callable[[Any], "list[Section]"]


# --------------------------------------------------------------------------- #
# 数据模型                                                                      #
# --------------------------------------------------------------------------- #
class Section(BaseModel):
    """结构化文档区块：一个标题及其直接归属的段落 (Req 5.2)。"""

    model_config = ConfigDict(extra="forbid")

    title: str = Field(
        default="",
        description="标题文本；顶层无标题正文使用空串占位。",
    )
    level: int = Field(
        default=0,
        ge=0,
        description="标题层级：0 表示无标题/正文根，1=H1，2=H2 ……",
    )
    paragraphs: list[str] = Field(
        default_factory=list,
        description="该标题直接归属的段落集合，保留原文顺序。",
    )


class FileMeta(BaseModel):
    """文件元数据：至少包含来源文件标识、文件名与文件格式 (Req 5.4)。"""

    model_config = ConfigDict(extra="forbid")

    source_file_id: str = Field(..., min_length=1, description="来源文件标识。")
    file_name: str = Field(..., min_length=1, description="原始文件名。")
    file_format: str = Field(
        ...,
        min_length=1,
        description="文件格式/扩展名，统一为小写且不含前导点。",
    )
    extra: dict[str, Any] = Field(
        default_factory=dict,
        description="可选的附加元数据（不影响下游必备字段）。",
    )


class ParseResult(BaseModel):
    """Loader 解析结果：统一纯文本 + 结构化区块 + 文件元数据。"""

    model_config = ConfigDict(extra="forbid")

    text: str = Field(..., description="统一纯文本内容表示 (Req 5.1)。")
    sections: list[Section] = Field(
        default_factory=list,
        description="保留标题层级顺序与段落归属的区块集合 (Req 5.2)。",
    )
    meta: FileMeta = Field(..., description="文件元数据 (Req 5.4)。")


# --------------------------------------------------------------------------- #
# 纯函数式解析辅助                                                              #
# --------------------------------------------------------------------------- #
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BLANK_LINE_RE = re.compile(r"\n[ \t]*\n")
_WS_RUN_RE = re.compile(r"\s+")
_HEADING_TAGS = {f"h{i}": i for i in range(1, 7)}


def _collapse_ws(text: str) -> str:
    """折叠连续空白为单个空格并去除首尾空白。"""
    return _WS_RUN_RE.sub(" ", text).strip()


def _decode(content: Any) -> str:
    """将原始内容解码为 ``str``；非 UTF-8 字节视为解析失败 (Req 5.3)。"""
    if isinstance(content, str):
        return content
    if isinstance(content, (bytes, bytearray)):
        try:
            return bytes(content).decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ParseFailedError(f"文件解码失败（非 UTF-8 编码）: {exc}") from exc
    raise ParseFailedError(f"不支持的内容类型: {type(content).__name__}")


def _split_into_paragraphs(text: str) -> list[str]:
    """按空行将纯文本切分为段落，单段落内多行折叠为一段。"""
    paragraphs: list[str] = []
    for block in _BLANK_LINE_RE.split(text):
        joined = " ".join(line.strip() for line in block.splitlines() if line.strip())
        joined = joined.strip()
        if joined:
            paragraphs.append(joined)
    return paragraphs


def _finalize_section(sections: list[Section], section: Section) -> None:
    """将区块收尾：仅当其为真正标题（level>0）或含段落时才纳入结果。"""
    if section.level > 0 or section.paragraphs:
        sections.append(section)


def _parse_txt(content: Any) -> list[Section]:
    """纯文本：单个 level-0 区块，段落按空行切分。"""
    text = _decode(content)
    paragraphs = _split_into_paragraphs(text)
    if not paragraphs:
        return []
    return [Section(title="", level=0, paragraphs=paragraphs)]


def _parse_markdown(content: Any) -> list[Section]:
    """Markdown：按 ATX 标题（``#``..``######``）保留层级与段落归属 (Req 5.2)。"""
    text = _decode(content)
    sections: list[Section] = []
    current = Section(title="", level=0, paragraphs=[])
    buffer: list[str] = []

    def flush_paragraph() -> None:
        if buffer:
            para = " ".join(line.strip() for line in buffer if line.strip()).strip()
            if para:
                current.paragraphs.append(para)
            buffer.clear()

    for line in text.splitlines():
        match = _HEADING_RE.match(line)
        if match is not None:
            flush_paragraph()
            _finalize_section(sections, current)
            current = Section(
                title=match.group(2).strip(),
                level=len(match.group(1)),
                paragraphs=[],
            )
        elif line.strip() == "":
            flush_paragraph()
        else:
            buffer.append(line)

    flush_paragraph()
    _finalize_section(sections, current)
    return sections


class _HtmlSectionExtractor(HTMLParser):
    """从 HTML 中提取 ``<h1>``..``<h6>`` 标题与 ``<p>``/松散文本段落。"""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.sections: list[Section] = []
        self._current = Section(title="", level=0, paragraphs=[])
        self._mode: str | None = None  # None | "heading" | "paragraph"
        self._heading_level = 0
        self._buf: list[str] = []
        self._loose: list[str] = []

    def _flush_loose(self) -> None:
        if self._loose:
            para = _collapse_ws("".join(self._loose))
            if para:
                self._current.paragraphs.append(para)
            self._loose.clear()

    def handle_starttag(self, tag: str, attrs: Any) -> None:
        tag = tag.lower()
        if tag in _HEADING_TAGS:
            self._flush_loose()
            self._mode = "heading"
            self._heading_level = _HEADING_TAGS[tag]
            self._buf = []
        elif tag == "p":
            self._flush_loose()
            self._mode = "paragraph"
            self._buf = []

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in _HEADING_TAGS and self._mode == "heading":
            title = _collapse_ws("".join(self._buf))
            _finalize_section(self.sections, self._current)
            self._current = Section(
                title=title, level=self._heading_level, paragraphs=[]
            )
            self._mode = None
            self._buf = []
        elif tag == "p" and self._mode == "paragraph":
            para = _collapse_ws("".join(self._buf))
            if para:
                self._current.paragraphs.append(para)
            self._mode = None
            self._buf = []

    def handle_data(self, data: str) -> None:
        if self._mode in ("heading", "paragraph"):
            self._buf.append(data)
        elif data.strip():
            self._loose.append(data)

    def result(self) -> list[Section]:
        """收尾：刷新松散文本并归档最后一个区块。"""
        self._flush_loose()
        _finalize_section(self.sections, self._current)
        return self.sections


def _parse_html(content: Any) -> list[Section]:
    """HTML：提取标题层级与段落 (Req 5.2)。"""
    text = _decode(content)
    extractor = _HtmlSectionExtractor()
    extractor.feed(text)
    extractor.close()
    return extractor.result()


def _parse_pdf(content: Any) -> list[Section]:
    """PDF：惰性依赖 ``pypdf``；缺失或解析失败时抛出 ParseFailedError。"""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - 取决于可选依赖是否安装
        raise ParseFailedError("缺少 PDF 解析依赖（pypdf 未安装）") from exc

    raw = content if isinstance(content, (bytes, bytearray)) else _decode(content).encode("utf-8")
    reader = PdfReader(BytesIO(bytes(raw)))
    texts: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            texts.append(page_text)
    paragraphs = _split_into_paragraphs("\n\n".join(texts))
    if not paragraphs:
        return []
    return [Section(title="", level=0, paragraphs=paragraphs)]


def _parse_docx(content: Any) -> list[Section]:
    """DOCX：惰性依赖 ``python-docx``；按 Heading 样式还原标题层级。"""
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - 取决于可选依赖是否安装
        raise ParseFailedError("缺少 DOCX 解析依赖（python-docx 未安装）") from exc

    raw = content if isinstance(content, (bytes, bytearray)) else _decode(content).encode("utf-8")
    document = docx.Document(BytesIO(bytes(raw)))
    sections: list[Section] = []
    current = Section(title="", level=0, paragraphs=[])
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = (para.style.name if para.style else "") or ""
        lowered = style_name.lower()
        if lowered.startswith("heading"):
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = int(digits) if digits else 1
            _finalize_section(sections, current)
            current = Section(title=text, level=level, paragraphs=[])
        else:
            current.paragraphs.append(text)
    _finalize_section(sections, current)
    return sections


#: 内置解析器注册表。pdf/docx 为惰性可选实现，缺依赖时优雅失败。
_BUILTIN_PARSERS: dict[str, ParserFn] = {
    "txt": _parse_txt,
    "text": _parse_txt,
    "md": _parse_markdown,
    "markdown": _parse_markdown,
    "html": _parse_html,
    "htm": _parse_html,
    "pdf": _parse_pdf,
    "docx": _parse_docx,
}


# --------------------------------------------------------------------------- #
# Loader 主体                                                                   #
# --------------------------------------------------------------------------- #
class Loader:
    """文档加载器：解析受支持类型文件为 :class:`ParseResult` (Req 5)。"""

    def __init__(
        self,
        *,
        supported_types: "frozenset[str] | set[str] | None" = None,
        parse_timeout_seconds: float | None = None,
    ) -> None:
        if supported_types is None:
            supported_types = get_settings().supported_document_types
        self._supported_types: set[str] = {
            self._normalize_format(item) for item in supported_types
        }
        self._timeout: float = (
            parse_timeout_seconds
            if parse_timeout_seconds is not None
            else DEFAULT_PARSE_TIMEOUT_SECONDS
        )
        self._parsers: dict[str, ParserFn] = dict(_BUILTIN_PARSERS)

    # ---------------------------- 公共 API -------------------------------- #
    def register_parser(
        self,
        file_format: str,
        parser: ParserFn,
        *,
        supported: bool = True,
    ) -> None:
        """注册/覆盖某文件格式的解析器（扩展点，亦便于测试注入）。

        Args:
            file_format: 文件格式（大小写与前导点均会被归一化）。
            parser: 解析器函数，接收原始内容并返回 ``list[Section]``。
            supported: 是否同时将该格式纳入受支持类型集合（默认 True）。
        """
        fmt = self._normalize_format(file_format)
        self._parsers[fmt] = parser
        if supported:
            self._supported_types.add(fmt)

    def parse(
        self,
        *,
        content: bytes | str,
        file_name: str,
        file_format: str,
        source_file_id: str | None = None,
        ingest_task: IngestTask | None = None,
    ) -> ParseResult:
        """将文件解析为 :class:`ParseResult`。

        成功时返回带元数据的解析结果（Req 5.1/5.2/5.4）。解析失败（含不支持的
        类型/解码失败，Req 5.3）、解析超时（Req 5.5）或内容为空（Req 5.6）时，
        终止处理、不移交 Transformer，将 ``ingest_task`` 标记为 ``FAILED`` 并
        抛出相应错误。

        Raises:
            ParseFailedError: 无法解析为纯文本（含不支持类型、解码失败）。
            ParseTimeoutError: 解析耗时达到配置上限仍未完成。
            EmptyContentError: 解析完成但提取到的纯文本为空。
        """
        try:
            fmt = self._normalize_format(file_format)
            parser = self._get_parser(fmt)
            try:
                sections = self._run_with_timeout(lambda: parser(content))
            except (ParseFailedError, ParseTimeoutError):
                raise
            except Exception as exc:  # 兜底：任意解析异常归类为解析失败 (Req 5.3)
                raise ParseFailedError(
                    f"文件 {file_name} 解析失败: {exc}"
                ) from exc

            text = self._build_text(sections)
            if not text.strip():
                raise EmptyContentError(f"文件 {file_name} 解析后内容为空")
        except (ParseFailedError, ParseTimeoutError, EmptyContentError) as exc:
            self._mark_failed(ingest_task, exc.message)
            raise

        meta = FileMeta(
            source_file_id=source_file_id or str(uuid.uuid4()),
            file_name=file_name,
            file_format=fmt,
        )
        return ParseResult(text=text, sections=sections, meta=meta)

    # ---------------------------- 内部实现 -------------------------------- #
    @staticmethod
    def _normalize_format(value: Any) -> str:
        """归一化文件格式：去首尾空白、转小写、去前导点。"""
        return str(value).strip().lower().lstrip(".")

    def _get_parser(self, fmt: str) -> ParserFn:
        """按格式取解析器；不支持/未注册均视为解析失败 (Req 5.3)。"""
        if fmt not in self._supported_types:
            supported = ", ".join(sorted(self._supported_types))
            raise ParseFailedError(
                f"不支持的文档类型: {fmt!r}；受支持类型: [{supported}]"
            )
        parser = self._parsers.get(fmt)
        if parser is None:
            raise ParseFailedError(f"未注册 {fmt!r} 类型的解析器")
        return parser

    def _run_with_timeout(self, fn: Callable[[], list[Section]]) -> list[Section]:
        """在工作线程中执行解析并施加超时守卫 (Req 5.5)。"""
        executor = ThreadPoolExecutor(max_workers=1)
        future = executor.submit(fn)
        try:
            return future.result(timeout=self._timeout)
        except FuturesTimeoutError as exc:
            raise ParseTimeoutError(
                f"文档解析超时（超过 {self._timeout} 秒）"
            ) from exc
        finally:
            # 不阻塞等待可能仍在运行的工作线程。
            executor.shutdown(wait=False)

    @staticmethod
    def _build_text(sections: list[Section]) -> str:
        """将区块还原为统一纯文本：标题与段落按顺序拼接 (Req 5.1)。"""
        parts: list[str] = []
        for section in sections:
            if section.title:
                parts.append(section.title)
            parts.extend(section.paragraphs)
        return "\n\n".join(parts)

    @staticmethod
    def _mark_failed(ingest_task: IngestTask | None, reason: str) -> None:
        """将入库任务标记为 FAILED 并附失败原因 (Req 5.3/5.5/5.6)。

        ``IngestTask`` 启用 ``validate_assignment``，且 ``status=FAILED`` 要求
        ``failure_reason`` 非空，因此必须先写 ``failure_reason`` 再写 ``status``。
        """
        if ingest_task is None:
            return
        ingest_task.failure_reason = reason
        ingest_task.status = IngestTaskStatus.FAILED
